import logging
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.common.exceptions import ElementClickInterceptedException, NoSuchElementException
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from persian_tools import digits
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from persiantools.jdatetime import JalaliDate
import os
import time
import urllib.request
from selenium.common.exceptions import StaleElementReferenceException
from pathlib import Path


def create_driver():
    chrome_options = Options()
    # chrome_options.add_argument('--headless')
    # chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_argument('--ignore-certificate-errors')
    chrome_options.add_argument('--ignore-ssl-errors')
    chrome_options.add_argument("--disable-popup-blocking")
    chrome_options.add_argument("--start-maximized")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-extensions")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-infobars")
    chrome_options.add_argument("--disable-browser-side-navigation")
    chrome_options.add_argument("--disable-images")
    # chrome_options.add_argument("user-data-dir=./cache")
    return webdriver.Chrome(options=chrome_options)

if __name__ == "__main__":
    driver = create_driver()

# logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')


# logger = logging.getLogger('selenium')
t_prices = []
d_prices = []

def locate_element_with_retry(driver, by, value, retries=3):
    for attempt in range(retries):
        try:
            return driver.find_element(by, value)
        except StaleElementReferenceException:
            if attempt < retries - 1:
                time.sleep(1)  # Optional: short delay before retrying
            else:
                raise

def locate_elements_with_retry(driver, by, value, retries=3):
    for attempt in range(retries):
        try:
            return driver.find_elements(by, value)
        except StaleElementReferenceException:
            if attempt < retries - 1:
                time.sleep(1)  # Optional: short delay before retrying
            else:
                raise


digi_urls = {
    "iPhone 16 CH-128-8" : "https://www.digikala.com/product/dkp-17986495/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%D9%84-iphone-16-ch-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=17986495&variant_id=69248083",  
    "iPhone 16 Pro Max ZAA-256-8" : "https://www.digikala.com/product/dkp-17986675/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%D9%84-iphone-16-pro-max-zaa-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=17986675&variant_id=81495686",    
    "iPhone 17 CH-256-8" : "https://www.digikala.com/product/dkp-20127115/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%D9%84-iphone-17-ch-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%86%D8%A7%D8%AA-%D8%A7%DA%A9%D8%AA%DB%8C%D9%88/?product_id=20127115&variant_id=75094889",    
    "iPhone 17 Pro ZAA-256-12" : "https://www.digikala.com/product/dkp-20125739/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%D9%84-iphone-17-pro-zaa-%D8%AA%DA%A9-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-esim-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%86%D8%A7%D8%AA-%D8%A7%DA%A9%D8%AA%DB%8C%D9%88/?product_id=20125739&variant_id=80856719",    
    "iPhone 17 Pro Max ZAA-256-12" : "https://www.digikala.com/product/dkp-20481188/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%D9%84-iphone-17-pro-max-zaa-%D8%AA%DA%A9-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-esim-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%86%D8%A7%D8%AA-%D8%A7%DA%A9%D8%AA%DB%8C%D9%88/?product_id=20481188&variant_id=73071289",
    "iPhone 17 Pro Max ZAA-512-12" : "https://www.digikala.com/product/dkp-20125788/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%D9%84-iphone-17-pro-max-zaa-%D8%AA%DA%A9-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-esim-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%86%D8%A7%D8%AA-%D8%A7%DA%A9%D8%AA%DB%8C%D9%88/?product_id=20125788&variant_id=73071326",


    # -----------------------------------------------------------------------------------------------------------------------------------------------------------
        
    "A07-64-4": "https://www.digikala.com/product/dkp-20109389/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a07-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-64-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20109389&variant_id=74547381",
    "A07-128-4": "https://www.digikala.com/product/dkp-20109389/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a07-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-64-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20109389&variant_id=74547381",
    "A07-128-6": "https://www.digikala.com/product/dkp-20110013/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a07-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-6-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20110013&variant_id=74547560",
    "A17-128-4": "https://www.digikala.com/product/dkp-20490694/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a17-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85/?product_id=20490694&variant_id=74561057",
    "A17-128-6": "https://www.digikala.com/product/dkp-20490714/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a17-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-6-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85/?product_id=20490714&variant_id=74561179",
    "A17-256-8": "https://www.digikala.com/product/dkp-20490734/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a17-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85/?product_id=20490734&variant_id=74561129",
    "A27-128-6": "",
    "A27-256-8": "",
    "A37-128-8": "https://www.digikala.com/product/dkp-21670148/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a37-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=21670148&variant_id=81408274",
    "A37-256-8": "https://www.digikala.com/product/dkp-21670150/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a37-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=21670150&variant_id=81170454",
    "A57-128-8": "https://www.digikala.com/product/dkp-21660233/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a57-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=21660233&variant_id=81408457",
    "A57-256-8": "https://www.digikala.com/product/dkp-21660239/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a57-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85/?product_id=21660239&variant_id=80582900",
    "A57-512-12": "https://www.digikala.com/product/dkp-21660243/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a57-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85/?product_id=21660243&variant_id=81389345",
    "S26 FE-256-8": "https://www.digikala.com/product/dkp-20135132/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-s25-fe-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85/?product_id=20135132&variant_id=81666257",
    "S26 Ultra-256-12": "https://www.digikala.com/product/dkp-21462684/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-s26-ultra-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85/?product_id=21462684&variant_id=77823184",
    "S26 Ultra-512-12": "https://www.digikala.com/product/dkp-21462685/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-s26-ultra-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85/?product_id=21462685&variant_id=79926941",
    
    # -----------------------------------------------------------------------------------------------------------------------------------------------------------
    
    "Redmi A3-128-4": "https://www.digikala.com/product/dkp-14503590/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-a3-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=14503590&variant_id=73760963",
    "Redmi A3 Pro-128-4": "https://www.digikala.com/product/dkp-20897081/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-a3-pro-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20897081&variant_id=80665349",
    "Poco C71-64-3": "https://www.digikala.com/product/dkp-19357767/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-poco-c71-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-64-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-3-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=19357767&variant_id=68134513",
    "Poco C71-128-4": "https://www.digikala.com/product/dkp-19347047/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-poco-c71-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=19347047&variant_id=73143292",
    "Redmi A5-128-4": "https://www.digikala.com/product/dkp-19742420/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-a5-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=19742420&variant_id=79162962",
    "Poco C81 Pro-128-4": "https://www.digikala.com/product/dkp-21865173/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-poco-c81-pro-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=21865173&variant_id=80438332",
    "Redmi 15C-128-4": "https://www.digikala.com/product/dkp-20434083/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-15c-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20434083&variant_id=74477588",
    "Redmi 15C-256-8": "https://www.digikala.com/product/dkp-20207335/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-15c-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20207335&variant_id=73240644",
    "Poco C85-128-6": "https://www.digikala.com/product/dkp-20434028/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-poco-c85-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-6-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20434028&variant_id=75670146",
    "Poco C85-256-8": "https://www.digikala.com/product/dkp-20433985/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-poco-c85-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20433985&variant_id=80775712",
    "Redmi 15C-256-8": "https://www.digikala.com/product/dkp-20207335/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-15c-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20207335&variant_id=73240644",
    "Redmi 13x-256-8": "https://www.digikala.com/product/dkp-19038990/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-13x-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=19038990&variant_id=73240736",
    "Redmi 15-256-8": "https://www.digikala.com/product/dkp-20607579/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-15-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=20607579&variant_id=73115559",
    "Redmi Note 14 4g-256-8": "https://www.digikala.com/product/dkp-21408633/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-14-5g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-%D9%BE%DA%A9-%DA%86%DB%8C%D9%86-%D9%88-%D8%B1%D8%A7%D9%85-%DA%AF%D9%84%D9%88%D8%A8%D8%A7%D9%84/?product_id=21408633&variant_id=76884302",
    "Redmi Note 15-256-8": "https://www.digikala.com/product/dkp-21190850/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-15-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=21190850&variant_id=81850548",
    "Redmi Note 15 pro-256-8": "https://www.digikala.com/product/dkp-21191058/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-15-pro-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=21191058&variant_id=77863195",
    "Redmi Note 15 pro-512-12": "https://www.digikala.com/product/dkp-21191071/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-15-pro-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=21191071&variant_id=77862828",
    "Redmi Note 15 pro plus-512-12": "https://www.digikala.com/product/dkp-21185646/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-15-pro-plus-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D9%88-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA/?product_id=21185646&variant_id=77862861",
}


techno_urls = {
    "iPhone 16 CH-128-4" : "https://www.technolife.com/product-69610/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%D9%84-iphone-16-ch-a-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D8%AC%DB%8C%D8%B3%D8%AA%D8%B1-%D8%B4%D8%AF%D9%87",
    "iPhone 16 Pro Max ZAA-256-8" : "https://www.technolife.com/product-59177/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%D9%84-iphone-16-pro-max-za-a-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D8%AC%DB%8C%D8%B3%D8%AA%D8%B1-%D8%B4%D8%AF%D9%87",  
    "iPhone 17 CH-256-8" : "https://www.technolife.com/product-165487/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-%D9%85%D8%AF%E2%80%8C%D9%84-iphone-17-ch-a-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---not-active-%D8%B1%D8%AC%DB%8C%D8%B3%D8%AA%D8%B1-%D8%B4%D8%AF%D9%87",
    "iPhone 17 Pro ZAA-256-12" : "https://www.technolife.com/product-149692/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-iphone-17-pro-za-a-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---not-active-%D8%B1%D8%AC%DB%8C%D8%B3%D8%AA%D8%B1-%D8%B4%D8%AF%D9%87",
    "iPhone 17 Pro Max ZAA-256-12" : "https://www.technolife.com/product-149697/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-iphone-17-pro-max-za-a-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---not-active-%D8%B1%D8%AC%DB%8C%D8%B3%D8%AA%D8%B1-%D8%B4%D8%AF%D9%87",
    "iPhone 17 Pro Max ZAA-512-12" : "https://www.technolife.com/product-165862/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%A7%D9%BE%D9%84-iphone-17-pro-max-za-a-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---not-active-%D8%B1%D8%AC%DB%8C%D8%B3%D8%AA%D8%B1-%D8%B4%D8%AF%D9%87",


    # -----------------------------------------------------------------------------------------------------------------------------------------------------------
        
    "A07-64-4": "https://www.technolife.com/product-149347/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a07-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-64-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "A07-128-4": "https://www.technolife.com/product-149351/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a07-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "A07-128-6": "https://www.technolife.com/product-149352/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a07-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-6-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "A17-128-4": "https://www.technolife.com/product-167564/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a17-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85",
    "A17-128-6": "https://www.technolife.com/product-167567/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a17-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-6-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85",
    "A17-256-8": "https://www.technolife.com/product-167563/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-a17-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85-",
    "A27-128-6": "https://www.technolife.com/product-569542/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-a27-5g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-6-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85",
    "A27-128-8": "https://www.technolife.com/product-569543/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-a27-5g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85",
    "A37-128-8": "https://www.technolife.com/product-529048/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-a37-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85",
    "A37-256-8": "https://www.technolife.com/product-400669/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-a37-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85",
    "A57-128-8": "https://www.technolife.com/product-529046/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-a57-5g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85?query_id=4c1a186e-456b-46d0-8f96-58e7b7eaaa68&position=3",
    "A57-256-8": "https://www.technolife.com/product-315067/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-a57-5g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=4c1a186e-456b-46d0-8f96-58e7b7eaaa68&position=2",
    "A57-512-12": "https://www.technolife.com/product-391888/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-a57-5g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "S26 FE-256-8": "https://www.technolife.com/product-165914/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-%D9%85%D8%AF%D9%84-galaxy-s25-fe-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "S26 Ultra-256-12": "https://www.technolife.com/product-390330/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-s26-ultra-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85",
    "S26 Ultra-512-12": "https://www.technolife.com/product-78111/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B3%D8%A7%D9%85%D8%B3%D9%88%D9%86%DA%AF-galaxy-s25-ultra-5g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA---%D9%88%DB%8C%D8%AA%D9%86%D8%A7%D9%85",
    
    # -----------------------------------------------------------------------------------------------------------------------------------------------------------
    
    "Redmi A3-128-4": "https://www.technolife.com/product-34773/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-a3-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "Redmi A3 Pro-128-4": "https://www.technolife.com/product-71107/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-a3-pro-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "Poco C71-64-3": "https://www.technolife.com/product-116567/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-poco-c71-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-64-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-3-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=8F774A22F4A9A44C&position=2",
    "Poco C71-128-4": "https://www.technolife.com/product-116578/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-poco-c71-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=8F774A22F4A9A44C&position=1",
    "Redmi A5-128-4": "https://www.technolife.com/product-107905/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-a5-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=5CB9EBF46F823D63&position=1",
    "Poco C81 Pro-128-4": "https://www.technolife.com/product-559198/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D9%BE%D9%88%DA%A9%D9%88-%D9%85%D8%AF%D9%84-c81-pro-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "Redmi 15C-128-4": "https://www.technolife.com/product-165996/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-redmi-15c-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-4-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=07874fad-39d0-4037-a64e-fc7c7a82a3e1&position=2",
    "Redmi 15C-256-8": "https://www.technolife.com/product-166010/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-redmi-15c-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=07874fad-39d0-4037-a64e-fc7c7a82a3e1&position=1",
    "Poco C85-128-6": "https://www.technolife.com/product-166017/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D9%BE%D9%88%DA%A9%D9%88-%D9%85%D8%AF%D9%84-c85-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-128-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-6-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "Poco C85-256-8": "https://www.technolife.com/product-166041/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D9%BE%D9%88%DA%A9%D9%88-%D9%85%D8%AF%D9%84-c85-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=b4fb1863-7f45-445b-a12a-e2a31604a7f7&position=1",
    "Redmi 13x-256-8": "https://www.technolife.com/product-116168/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-13x-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=57795c81-354a-4df0-b3cc-671d36c9b467&position=1",
    "Redmi 15C-256-8": "https://www.technolife.com/product-166010/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-redmi-15c-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "Redmi 15-256-8": "https://www.technolife.com/product-173207/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%DB%8C%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-15-4g-%D8%AF%D9%88-%D8%B3%DB%8C%D9%85-%DA%A9%D8%A7%D8%B1%D8%AA-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA",
    "Redmi Note 14 4g-256-8": "https://www.technolife.com/product-76419/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-14-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=05FE66828BA32D4B&position=1",
    "Redmi Note 15-256-8": "https://www.technolife.com/product-314976/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-15-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=6fbed264-c081-4755-b6c3-70d22bb5ffef&position=1",
    "Redmi Note 15 pro-256-8": "https://www.technolife.com/product-301544/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-15-pro-4g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-8-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=6fbed264-c081-4755-b6c3-70d22bb5ffef&position=2",
    "Redmi Note 15 pro-512-12": "https://www.technolife.com/product-387549/%DA%AF%D9%88%E2%80%8C%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-15-pro-5g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-256-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=6fbed264-c081-4755-b6c3-70d22bb5ffef&position=3",
    "Redmi Note 15 pro plus-512-12": "https://www.technolife.com/product-387540/%DA%AF%D9%88%D8%B4%DB%8C-%D9%85%D9%88%D8%A8%D8%A7%D9%8A%D9%84-%D8%B4%DB%8C%D8%A7%D8%A6%D9%88%D9%85%DB%8C-%D9%85%D8%AF%D9%84-redmi-note-15-pro-plus-5g-%D8%B8%D8%B1%D9%81%DB%8C%D8%AA-512-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA-%D8%B1%D9%85-12-%DA%AF%DB%8C%DA%AF%D8%A7%D8%A8%D8%A7%DB%8C%D8%AA?query_id=6fbed264-c081-4755-b6c3-70d22bb5ffef&position=4",
}

xpath_for_black_techno = [
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[1]/div/p[contains(text() , "مشکی")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[2]/div/p[contains(text() , "مشکی")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[3]/div/p[contains(text() , "مشکی")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[4]/div/p[contains(text() , "مشکی")]'
]
xpath_for_darkblue = [
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[1]/div/p[contains(text() , "سرمه‌ای")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[2]/div/p[contains(text() , "سرمه‌ای")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[3]/div/p[contains(text() , "سرمه‌ای")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[4]/div/p[contains(text() , "سرمه‌ای")]'
]

xpath_for_white = [
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[1]/div/p[contains(text() , "سفید")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[2]/div/p[contains(text() , "سفید")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[3]/div/p[contains(text() , "سفید")]',
    '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[1]/div/div[3]/div/div[2]/div/div/div/div/div[4]/div/p[contains(text() , "سفید")]'
]


xpath_for_price_techno = {
    '1': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[4]/div/div/div/p',
    '2': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div[2]/div[3]/div[2]/div[2]/div/div/p[2]',
    '3': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[2]/div[2]/div/div/p',
    '4': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div/div[3]/div[2]/div/div/div/p',
    '5': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div[2]/div[3]/div[2]/div/div/div/p[2]',
    '6': '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[1]/div/div[2]/div[3]/div[4]/div/div/div/p[2]'
}

def check_internet_connection():
    # testing the connection by pinging google
    try:
        urllib.request.urlopen('https://www.google.com/', timeout=5)
        return True
    except:
        return False

def wait_for_connection(max_retries=10, retry_delay=10):
    # waiting for user to reconnect the connection
    retries = 0
    while retries < max_retries:
        if check_internet_connection():
            if retries > 0:
                print("Internet is connected....")
            return True
        else:
            retries += 1
            print(f"No internet connection. Retrying in {retry_delay} seconds... ({retries}/{max_retries})")
            time.sleep(retry_delay)

    print("Failed to reconnect after multiple attempts.")
    return False



if len(digi_urls) != len(techno_urls):
    raise Exception("The urls for technolife and digikala are diffrent")
else:
    urls_len = len(digi_urls)
    phone_models = []
    for key in digi_urls.keys():
        phone_models.append(key)
    # d_pbar = tqdm(total=urls_len)
    # t_pbar = tqdm(total=urls_len)
    # c_pbar = tqdm(total=100)
    # m_pbar = tqdm(total=3)

def deny(btn):
    try:
        # Wait for the 'deny' button to appear
        deny_btn = WebDriverWait(driver, 15).until(
            EC.presence_of_element_located((By.ID, "webpush-onsite"))
        )
        iframe = driver.find_element(By.ID, 'webpush-onsite')
        driver.switch_to.frame(iframe)
        
        # Try clicking the deny button
        try:
            deny_btn.click()
        except Exception as e:
            # if isinstance(e, ElementClickInterceptedException):
            #     logger.debug("Debug : ElementClickInterceptedException")
            # else:
            #     logger.debug(f"Debug : Exception occurred - {type(e).__name__}")
            
            try:
                # Try clicking with XPATH as fallback
                driver.find_element(By.XPATH, '//*[@id="deny"]').click()
            except Exception as inner_e:
                # logger.debug(f"Debug : Failed to click deny button - {type(inner_e).__name__}")
                t_prices.append('//')
                print('*/')
                driver.implicitly_wait(500)
                return 1
        else:
            # Default action if deny button is clicked successfully
            btn.click()
    except TimeoutException:
        # logger.debug("Debug : DenyButtonNotFound [In_Time]")
        t_prices.append('//')
        print('/*')
        return 1
    finally:
        # Switch back to the main content in all cases
        driver.switch_to.default_content()

def click_color(driver):
    # checking for the colors available
    try:
        colors_container = driver.find_element(
            By.CSS_SELECTOR, 
            ".gap-5 > div:nth-child(1) > div:nth-child(2) > div:nth-child(1)"
        )

        child_divs = colors_container.find_elements(By.CSS_SELECTOR, "div[id]")

        numeric_divs = [
            div for div in child_divs 
            if div.get_attribute("id") and div.get_attribute("id").isdigit()
        ]

        if not numeric_divs:
            print("No numeric IDs found")
            return False
        else:
            lowest_div = min(numeric_divs, key=lambda d: int(d.get_attribute("id")))

            svg_to_click = lowest_div.find_element(By.CSS_SELECTOR, "svg")
            svg_to_click.click()
            
            rang = "color"
    except NoSuchElementException:
        return False

    return rang

def normalize_price(price_text):
    return digits.convert_to_en(price_text)
    
def digi_scrape():
    for model , url in digi_urls.items():
        out_off_stock = True
        rang = False

        if url == r"": 
            d_prices.append("**")
            print(model , end="---**")
            continue
        
        if not wait_for_connection(max_retries=10, retry_delay=10):
            print("Could not establish connection. Exiting program.")
            return False
        else:
            driver.get(url)

        try:
            product_title = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "[data-testid='pdp-title']"))
               )      
            

            try:
                # driver.find_element(By.XPATH , '//*[@id="__next"]/div[1]/div[3]/div[3]/div[2]/div[2]/div[2]/div[2]/div[4]/div/div/div/button/div[2]/div')
                driver.find_element(By.XPATH , '//*[@id="__next"]/div[1]/div[3]/div[3]/div[2]/div[2]/div[2]/div[1]/div/h1/span')

            except NoSuchElementException:
                out_off_stock = False
            else:
                print(f"{model} **")
                d_prices.append('**')
                continue

            if click_color(driver):
                rang = "color"
            else:
                out_off_stock= True

            if rang:
                print(model , rang, end=" ")
            else:
                print(model , end=" ")
            
            try:
                price_no_discount = driver.find_element(By.CSS_SELECTOR , "div[data-theme-animation='price-container'] [data-testid='price-no-discount']")
                if 'line-trough' in price_no_discount.get_attribute("class"):
                    final_price_list = driver.find_element(By.CSS_SELECTOR , "div[data-theme-animation='price-container'] [data-testid='price-final']")

                    price = final_price_list
                else:
                    price = price_no_discount
            except NoSuchElementException:
                try:
                    final_price_list = driver.find_element(By.CSS_SELECTOR , "div[data-theme-animation='price-container'] [data-testid='price-final']")

                    price = final_price_list
                except NoSuchElementException:
                    d_prices.append("//")
                    print('//')
            

            if out_off_stock == False:
                if isinstance(price , str):
                    d_prices.append(price)
                    print(price)
                else:
                    final = normalize_price(price.text)
                    d_prices.append(final)
                    print(final)
        
        except TimeoutException:
            print(f"Failed to find the title for {url} within the given time.")
            d_prices.append('//')

        continue
        # d_pbar.update(1)


percent = 100 / len(techno_urls)

# loading the page 
def techno_scrape():
    for model , url in techno_urls.items():

        if url == r"": 
            out_off_stock = True
            t_prices.append("**")
            print(model , end="---**")
            continue
    
        if not wait_for_connection(max_retries=10, retry_delay=10):
            print("Could not establish connection. Exiting program.")
            return False
        else:
            driver.get(url)

        print(model , end="---")

        try:
            product_title = WebDriverWait(driver, 20).until(
                    EC.presence_of_element_located((By.ID, "pdp_name"))
                )     
            

            try:
                driver.find_element(By.XPATH , '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[2]/div/div/div/div/div/p[contains (text() , "ناموجود")]')
            except NoSuchElementException:
                pass
            else:
                t_prices.append("**")
                print('**')
                continue

            rang = 'N/A'
            out_off_stock = False
            price = "//"

            

            # cheking for the colors available
            try:
                driver.find_element(By.CSS_SELECTOR, "[style='background-color:#1a1a1a'").click()
            except NoSuchElementException:
                try:
                    driver.find_element(By.CSS_SELECTOR, "[style='background-color:#00009c']").click()
                except NoSuchElementException:
                    pass
                except ElementClickInterceptedException:
                    if deny(driver.find_element(By.CSS_SELECTOR, "[style='background-color:#00009c']")) == 1:
                        continue
                else:
                    rang = "DarkBlue"
            except ElementClickInterceptedException:
                if deny(driver.find_element(By.CSS_SELECTOR, "[style='background-color:#1a1a1a'")) == 1:
                    continue
            else:
                rang = "Black"

            # finding the price and scraping it
            for x in xpath_for_price_techno:
                try:
                    price = locate_element_with_retry(driver, By.XPATH , xpath_for_price_techno[x]) # driver.find_element(By.XPATH , xpath_for_price_techno[x])
                except NoSuchElementException:
                    pass
                else:
                    break
        


            if out_off_stock == False:
                if isinstance(price, str):
                    t_prices.append(price)
                    print(price)
                else:
                    t_prices.append(price.text)
                    print(price.text)
            
        except TimeoutException:
                print(f"Failed to find the title for {model} within the given time.")
                t_prices.append('//')


        continue
        # t_pbar.update(1)




def create_document(phone_models, d_prices, t_prices, output_path):
    # creating the document and the row
    document = Document()
    table = document.add_table(rows=1, cols=3)

    # giving the style values

    style = document.styles['Normal']
    table.style = 'Table Grid'
    style.font.name = "Calibri" # type: ignore
    style.font.size = Pt(20) # type: ignore

    # The array of Phone model names , digikala prices , technolife prices


    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('phone').bold = True
    hdr_cells[1].text = 'Digikala'
    hdr_cells[2].text = 'Technolife'

    for phone, digi_price, techno_price in zip(
        phone_models,
        d_prices,
        t_prices
    ):
        row_cells = table.add_row().cells
        row_cells[0].paragraphs[0].add_run(phone).bold = True
        row_cells[1].text = digi_price
        row_cells[2].text = techno_price


    # document.save(f"{today_date}.docx")
    # c_pbar.update(50)

    # Save temporary DOCX
    doc_file = str(output_path.with_suffix(".docx"))
    document.save(doc_file)

    # Convert DOCX to PDF
    pdf_file = str(output_path.with_suffix(".pdf"))
    convert(doc_file, pdf_file)

    os.remove(doc_file)

def single_digi_scrape(model):
    out_off_stock = True
    rang = False
    
    if not wait_for_connection(max_retries=10, retry_delay=10):
        print("Could not establish connection. Exiting program.")
        return False
    else:
        driver.get(digi_urls[model])



    try:
        product_title = WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((By.CSS_SELECTOR, "[data-testid='pdp-title']"))
            )      
        

        try:
            driver.find_element(By.XPATH , '//*[@id="__next"]/div[1]/div[3]/div[3]/div[2]/div[2]/div[2]/div[2]/div[4]/div/div/div/button/div[2]/div')
        except NoSuchElementException:
            out_off_stock = False
        else:
            print(f"{model} **")

        # cheking for the colors available
        try:
            driver.find_element(By.CSS_SELECTOR, "[style='background: rgb(0, 33, 113);']").click()
        except NoSuchElementException:
            try:
                driver.find_element(By.CSS_SELECTOR, "[style='background: rgb(33, 33, 33);']").click()
            except NoSuchElementException:
                pass
            else:
                rang = "Dark Blue"
        else:
            rang = "Black"

        if rang:
            print(model , rang, end=" ")
        else:
            print(model , end=" ")
        
        try:
            price = driver.find_element(By.CSS_SELECTOR , '[data-testid="price-no-discount"]')
        except NoSuchElementException:
            try:
                price = driver.find_element(By.CSS_SELECTOR , '[data-testid="price-final"]')
            except NoSuchElementException:
                print('//')
        

        if out_off_stock == False:
            if isinstance(price , str):
                print(price)
                return price
            else:
                final = digits.convert_to_en(price.text)
                print(final)
                return final
    
    except TimeoutException:
        print(f"Failed to find the title for {model} within the given time.")

    # d_pbar.update(1)


# loading the page 
def single_techno_scrape(model):    
    if techno_urls[model] == r"https://www.google.com": 
        out_off_stock = True
        print(model , end="--- **")
        return

    if not wait_for_connection(max_retries=10, retry_delay=10):
        print("Could not establish connection. Exiting program.")
        return False
    else:
        driver.get(techno_urls[model])

    print(model , end="---")

    try:
        product_title = WebDriverWait(driver, 20).until(
                EC.presence_of_element_located((By.ID, "pdp_name"))
            )     
        

        try:
            out_off_stock = driver.find_element(By.XPATH , '//*[@id="__next"]/div[3]/main/div/div/article[1]/section[2]/div/div[2]/div/div/div/div/div/p[contains (text() , "ناموجود")]')
        except NoSuchElementException:
            pass
        else:
            t_prices.append("**")
            print('**')
            

        rang = 'N/A'
        out_off_stock = False
        price = "//"

        

        # cheking for the colors available
        try:
            black_btn = driver.find_element(By.CSS_SELECTOR, "[style='background-color:#1a1a1a'")
        except NoSuchElementException:
            try:
                dark_blue_btn = driver.find_element(By.CSS_SELECTOR, "[style='background-color:#00009c']")
            except NoSuchElementException:
                pass
            else:
                try:
                    dark_blue_btn.click()                    
                except ElementClickInterceptedException:
                    if deny(dark_blue_btn) == 1:
                        return
                else:
                    rang = "DarkBlue"          
        else:
            try:
                black_btn.click()
            except ElementClickInterceptedException:
                if deny(black_btn) == 1:
                    return
            else:
                rang = "Black"
                


        # finding the price and scraping it
        for x in xpath_for_price_techno:
            try:
                price = driver.find_element(By.XPATH , xpath_for_price_techno[x])
            except NoSuchElementException:
                pass
            else:
                break
                
        if rang:
            print(rang, end="")


        if out_off_stock == False:
            if isinstance(price, str):
                print(price)
                return price
            else:
                print(price.text)
                return price.text
        
    except TimeoutException:
            print(f"Failed to find the title for {model} within the given time.")


def list_gen():
    try:
        digi_start = time.time()
        digi_scrape()
        digi_end = time.time()
        digi_time = digi_end - digi_start
        print(f"Digi time = {digi_time}")

        techno_start = time.time()
        techno_scrape()
        techno_end = time.time()
        techno_time = techno_end - techno_start
        print(f"Techno time = {techno_time}")
    finally:
        driver.quit()


    today_date = str(JalaliDate.today())
    file_name = today_date[5:]
    output_dir = Path(today_date[:-3])

    if not output_dir.exists():
        output_dir.mkdir(parents=True)

    output_path = output_dir / file_name

    create_document(
        phone_models,
        d_prices,
        t_prices,
        output_path
    )

def single_model():
    model_to_scrape = input("Enter the phone model you want to scrape (e.g., A05-64-4): ")
    
    if model_to_scrape not in digi_urls or model_to_scrape not in techno_urls:
        raise ValueError(f"Model '{model_to_scrape}' not found in URLs dictionaries.")

    techno_start = time.time()
    single_techno_scrape(model_to_scrape)
    techno_end = time.time()
    techno_time = techno_end - techno_start
    print(f"Digi time = {techno_time}")
    # m_pbar.update(1)

    digi_start = time.time()
    single_digi_scrape(model_to_scrape)
    digi_end = time.time()

    digi_time = digi_end - digi_start
    print(f"Digi time = {digi_time}")

    driver.quit()
    
    return


def main():
    if not wait_for_connection(max_retries=10, retry_delay=10):
        print("Could not establish connection. Exiting program.")
    else:
        while True:
            user_input = input("Do you want to generate the price list for phones...?(Y/N)")
            driver = create_driver()
            if user_input == 'Y' or user_input == 'y':
                list_gen()
                break
            elif user_input == 'N' or user_input == 'n':
                single_model()
                break
            else:
                print("Invalid input. Please enter 'Y' or 'N' \nPlease Try again(Y/N)")


if __name__ == "__main__":
    main()
